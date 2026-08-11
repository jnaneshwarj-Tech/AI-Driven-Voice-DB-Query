import os
import sys
from datetime import datetime

# Import reportlab components
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, PageBreak
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

# Import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

TITLE = "AI & NLP Driven Smart Student Database Management System using MySQL"
SUBTITLE = "Phase-1 Project Report"
DEPT = "DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING"
COLLEGE = "GOVERNMENT ENGINEERING COLLEGE MOSALEHOSAHALLI, HASSAN"
UNIVERSITY = "VISVESVARAYA TECHNOLOGICAL UNIVERSITY, BELAGAVI"
DATE_STR = datetime.now().strftime("%B %Y")

def set_cell_background(cell, hex_color):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def build_docx(filename="Phase1_Project_Report.docx"):
    print("Generating DOCX...")
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles Setup
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Title Page
    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_univ = p_univ.add_run(UNIVERSITY.upper())
    run_univ.font.size = Pt(14)
    run_univ.font.bold = True
    run_univ.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    
    p_title_spacer = doc.add_paragraph()
    for _ in range(5):
        p_title_spacer.add_run("\n")
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(TITLE.upper())
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(SUBTITLE)
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    for _ in range(7):
        doc.add_paragraph()
        
    p_dept = doc.add_paragraph()
    p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_dept = p_dept.add_run(DEPT + "\n" + COLLEGE)
    run_dept.font.size = Pt(12)
    run_dept.font.bold = True
    run_dept.font.color.rgb = RGBColor(0x2A, 0x52, 0x98)
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(DATE_STR)
    run_date.font.size = Pt(11)
    run_date.font.bold = True
    
    doc.add_page_break()
    
    # ── Report Content Definitions ─────────────────────────────────────────────
    sections_content = [
        ("1. Abstract", 
         "The AI & NLP Driven Smart Student Database Management System is a state-of-the-art enterprise-grade "
         "academic ERP system that bridges the gap between natural human language and structured relational databases. "
         "By leveraging advanced NLP techniques, fuzzy phonetic matching algorithms, and RAG (Retrieval-Augmented Generation), "
         "the system translates text and voice queries directly into secure, highly optimized MySQL commands. "
         "It addresses crucial bottlenecks in institutional operations, including complex search ambiguities, academic status tracking, "
         "robust schema migrations, and full transactional recovery with a multi-layered undo snapshot mechanism. "
         "The application delivers premium institutional branding, highly resilient data parsing, and absolute security against "
         "SQL injection attacks, redefining digital student administration in accordance with modern pedagogical guidelines."),
        
        ("2. Introduction",
         "In modern academic settings, administrative tasks demand speed, precision, and data integrity. Traditional ERP solutions "
         "rely on complex menu trees, multiple filter selections, and precise SQL matching, resulting in steep learning curves "
         "and operational delays. The 'AI & NLP Driven Smart Student Database Management System' is developed to resolve "
         "these limitations. By integrating intuitive search with relational MySQL databases, it allows users—both tech-savvy "
         "administrators and non-technical staff—to command the system via plain text and spoken language. "
         "This Phase-1 report outlines the complete architectural layout, database structures, AI-powered parsing engines, "
         "7-layer fuzzy matching modules, and robust soft-delete pipelines that make this system exceptionally resilient, user-friendly, "
         "and highly responsive."),
        
        ("3. Problem Statement",
         "Academic databases suffer from severe operational friction: \n"
         "1. Traditional search engines fail to find records when users make typos or phonetic naming errors, causing "
         "redundant entry or administrative delays.\n"
         "2. When multiple students share identical or similar names, typical systems automatically merge, overwrite, "
         "or arbitrarily display the first match without letting the user select the appropriate student, creating data contamination.\n"
         "3. Bulk file uploads of spreadsheets containing varied column layouts require manually rearranging spreadsheets "
         "to fit rigid database constraints, leading to massive manual parsing effort.\n"
         "4. Accidentally deleted or updated records are immediately permanent in traditional relational setups, lacking "
         "an instant point-in-time recovery without restoring bulky daily database backups."),
        
        ("4. Objectives",
         "The core objectives of the system include:\n"
         "• Develop a multi-layer NLP-to-SQL search pipeline utilizing token-level fuzzy, trigram, and phonetic matching "
         "to enable highly lenient live typing suggestions (≥0.20 score) and exact submitted queries (≥0.70 threshold).\n"
         "• Implement an ambiguity detection workflow that intercepts multi-student matches with similar names, blocking "
         "unauthorized merges and prompting the user to select the specific student by USN.\n"
         "• Build an AI Column Mapping and Wide-to-Long Excel Parser that auto-detects headers, matches custom columns to "
         "canonical database columns via exact/partial matching, and unpivots wide-format semester metrics to normalized rows.\n"
         "• Formulate a VTU Semester Progression Logic to automatically calculate estimated semesters, academic years, lateral entry "
         "variances, and graduated status based on parsed USN structures.\n"
         "• Incorporate an undo-snapshot engine offering a 5-minute transaction recovery window for deletion, updates, "
         "and bulk uploads by leveraging soft-deleted JSON snapshots."),
        
        ("5. Existing vs Proposed System",
         "The current project replaces outdated, rigid, and slow search architectures with a state-of-the-art system. "
         "The differences are highlighted below:\n\n"
         "• SEARCH LOGIC: The existing system utilizes rigid exact SQL substring matching (e.g., LIKE '%name%'), which fails on typos. "
         "The proposed system uses a 7-layer NLP pipeline incorporating soundex, double metaphone, trigram, and Levenshtein metrics "
         "to resolve spelling and phonetic discrepancies.\n"
         "• BULK INGESTION: Traditional systems crash or reject spreadsheets if column names do not exactly match the database. "
         "The proposed system utilizes a rule-based AI Column Mapper that normalizes column headers and automatically parses wide-format "
         "semester matrices to canonical, normalized structures.\n"
         "• DELETION SAFETY: In typical architectures, database deletions are immediately permanent, risking complete data loss. "
         "The proposed system captures full point-in-time pre-deletion JSON snapshots, allowing instant restore via a unique UUID "
         "token within a 5-minute window.\n"
         "• BRANDED EXPORTS: Output reporting in standard tools returns generic unformatted tables. The proposed system employs a unified "
         "branded report generator in PDF, Excel, and CSV formats, including institutional headers and computer-filled signature sections."),
        
        ("6. Technologies Used",
         "• Backend Framework: Python FastAPI (v0.109.2) - High-performance ASGI framework enabling concurrent async API execution.\n"
         "• Frontend Architecture: React.js (v18) & Vite - Lightning-fast UI component loading and responsive design using TailwindCSS.\n"
         "• Database Engine: Oracle MySQL (v8.0) - Relational DB utilizing indexing, unique keys, and transactions.\n"
         "• Text & Voice Parsing: Web Speech API - Browser-native voice dictation converting oral requests directly into strings.\n"
         "• Data Modeling & Math: Pandas (v2.2.0) - High-speed dataframe manipulation used in file parsing.\n"
         "• Professional Reporting: ReportLab (v4.1.0) & OpenPyXL (v3.1.2) - Direct generation of branded PDF, Excel, and CSV exports.\n"
         "• Security & Authentication: Bcrypt (v4.1.2) & Python-Jose (v3.3.0) - Robust password hashing and JWT token handling."),
        
        ("7. System Architecture & Working Pipeline",
         "The operation follows a clear, logical progression: \n"
         "1. USER QUERY: Plain text input is received via the React input terminal or Web Speech API.\n"
         "2. NLP PROCESSING: Stop words are eliminated; the query is analyzed for intent (personal, academic, or full) and search terms are extracted.\n"
         "3. SQL QUERY GENERATION: RAG (Retrieval-Augmented Generation) feeds schema constraints and few-shot examples to the NVIDIA LLM API.\n"
         "4. SECURITY VALIDATION: The generated SQL is parsed by the security validator. Unsafe commands (DROP, ALTER, TRUNCATE) or destructive statements lacking a WHERE clause are blocked and logged in the security_logs table.\n"
         "5. DATABASE FETCH: Valid SELECT queries are fetched from the MySQL connection pool. If 0 records are returned, the system falls back to phonetic/fuzzy matching.\n"
         "6. DUAL-MODAL FEEDBACK: Ambiguities (multi-student name matches) are displayed via selection cards. Confirmed results are mapped and rendered in interactive tables, charts, or exported in high-quality formats."),
        
        ("8. Modules Description",
         "• NLP Parsing & Intent Detection: Strips stop words, determines query scope (academic vs. personal), and extracts student names/USNs.\n"
         "• 7-Layer Fuzzy Search Engine: Provides live, lenient, spelling-resilient suggestions during typing and strict confidence scoring upon submission.\n"
         "• AI File Upload & Column Mapper: Receives CSV/Excel files, normalizes columns, and unpivots wide semester columns into structured rows.\n"
         "• Unified Export System: Builds identical Report models to construct branded PDFs, Excel sheets, and CSVs with integrated signatures.\n"
         "• Undo/Snapshot Engine: Intercepts updates, deletes, and uploads to save pre-state JSON dumps for instant Point-In-Time rollback."),
        
        ("9. 7-Layer NLP Search Engine Logic",
         "Spelling-resilient matching is achieved using a robust 7-layer hybrid scoring and search pipeline:\n"
         "1. Exact String Match: Direct case-sensitive string equality checks (100% confidence).\n"
         "2. Normalized Exact Match: Lowercases input, strips accents via unicodedata, and trims special characters (97% confidence).\n"
         "3. Prefix & Token Prefix Match: Checks if candidate names or space-separated name tokens start with the query (e.g. 'rut' -> 'ruthik'). Utilizes a length-based coverage score ratio.\n"
         "4. Substring & Token Substring Match: Checks if normalized query exists inside candidate name tokens (e.g. 'anth' -> 'manjunath').\n"
         "5. Phonetic Encoding (Soundex & Double Metaphone): Generates standard 4-char English pronunciation keys (Soundex) and advanced multi-pronunciation hashes (Double Metaphone) via Jellyfish to identify words sounding identical despite spelling differences (e.g., 'Sudheer' -> 'Sudhir').\n"
         "6. Character Trigram Similarity: Extracts character-level trigrams and computes a similarity coefficient matching typos: 2.0 * len(intersection) / (len(trigrams_A) + len(trigrams_B)).\n"
         "7. Levenshtein & Jaro-Winkler String Distance: Calculates single-character insertions, deletions, substitutions, and transpositions to grade typographical proximity.\n"
         "• Live Suggestion Rule: Executes leniently with score threshold >= 0.20 for high-response keyboard inputs.\n"
         "• Submission Execution Rule: Applies strict score threshold >= 0.70 to ensure secure database write-back commands."),
        
        ("10. Smart Semester Progression Logic",
         "The system parses Visvesvaraya Technological University (VTU) USN structures (e.g. '1GC20CS042') using regex:\n"
         "• Course Duration & Year: First character digit defines course duration (e.g., 4 years). Characters 4-5 represent the short admission year (e.g. 20 -> 2020).\n"
         "• Student Type: Roll number digits (characters 8-10) are parsed to detect Lateral Entry (roll number >= 400). Regular students start from 1.\n"
         "• Progression Calculation: Compares admission year to current year. If the current month is >= July, the semester is calculated as (years_diff * 2 + 1) [Odd Semester]. Otherwise, it is (years_diff * 2) [Even Semester]. Lateral entry students automatically receive a +2 semester adjustment.\n"
         "• Status Mapping: If the estimated semester exceeds (course_duration * 2), status is set to GRADUATED; otherwise, ACTIVE."),
        
        ("11. File Ingestion & AI Mapping",
         "Spreadsheet upload is managed through a strict, zero-loss pipeline:\n"
         "1. Header Row Detection: Scans the first 20 rows of Excel/CSV sheets to locate the real header by identifying known keywords like USN, name, or GPA.\n"
         "2. AI Column Mapper: Maps custom headers to canonical database columns. For example, 'Reg Number' maps to 'usn', and 'Fathers Name' to 'father_name'.\n"
         "3. Wide-Format Unpivoting: Converts wide sheets (columns: sem_1_sgpa, sem_2_sgpa) into standardized normalized rows (Semester, SGPA).\n"
         "4. USN Validation & Merge: Checks USN structure using standard regex. If valid, records are merged (upserted); duplicate rows are identified and flagged."),
        
        ("12. Undo & Point-in-Time Recovery",
         "To guarantee complete transactional safety, a 5-minute Point-In-Time recovery window is enforced:\n"
         "• Deletion/Update Interception: Prior to executing any destructive command, the system runs a select query to fetch all target rows.\n"
         "• Snapshot Generation: Stores the full student profiles and academic marks as a serialized JSON dump in the `global_undo_snapshots` table, returning a unique UUID token.\n"
         "• Recovery Execution: Upon post-requesting the token to `/api/undo/restore/{token}`, the system deletes the newly modified data, reads the JSON snapshot, and re-inserts the exact original values."),
        
        ("13. MySQL Database Design",
         "The database is designed with optimized indexing and strict integrity constraints. Key tables include:\n"
         "• users: Manages credentials, roles (Admin/Staff), and preferences (theme) with custom ENUMs.\n"
         "• students: Central registry of personal fields (USN, name, DOB, father name, permanent address) with a UNIQUE USN constraint.\n"
         "• marks: Tracks academic data (SGPA, CGPA) linked to USN via foreign key cascading, protected by a unique index (uq_usn_sem).\n"
         "• global_undo_snapshots: Stores the pre-operation JSON data, operation type, actor, and status for point-in-time rollbacks.\n"
         "• query_history & security_logs: Maintain detailed auditing for natural queries, executed SQL statements, and intercepted injection attempts."),
        
        ("14. UI/UX Architecture",
         "The frontend is constructed using modular React component design:\n"
         "• Dashboard Workspace: Provides a single-page reactive shell hosting a collapsible sidebar and unified search hub.\n"
         "• Live Suggestion Panel: A highly interactive dropdown showing spelling matches with badges (EXACT, FUZZY, PHONETIC).\n"
         "• Interactive Data Tables: Generates responsive paginated grids with column sorting and search filtering.\n"
         "• Activity Feed & Rollback Panel: Allows administrators to review security logs and trigger instant undo rollbacks in real time."),
        
        ("15. Results & Implemented Features",
         "The Phase-1 implementation succeeds in delivering all core modules:\n"
         "✔ Verified 7-layer fuzzy phonetic search engine resolving severe typos.\n"
         "✔ Operational async text and voice-dictated query generation pipeline.\n"
         "✔ Zero-loss spreadsheet ingestion engine with automatic unpivoting.\n"
         "✔ Strict security validation blocking malicious AST injection patterns.\n"
         "✔ Verified 5-minute Point-In-Time Undo Recovery with JSON snapshots.\n"
         "✔ Institutional branded PDF, Excel, and CSV export modules."),
        
        ("16. Remaining Issues & Future Scope",
         "Current limitations and developmental plans for Phase-2 include:\n"
         "• Current Limitations: Local MySQL service configuration dependencies (requires manual administrative startup); single-instance file upload cache (susceptible to cache expiration on server restart).\n"
         "• Future Scope: Integration of vector embedding semantic search using Milvus/ChromaDB to handle complex intent mappings; implementation of Redis-based query caching to bypass LLM generation for identical queries; integration of automated WebSockets for multi-user real-time notification feeds."),
        
        ("17. Conclusion",
         "Phase-1 of the AI & NLP Driven Student Database Management System successfully establishes a robust, highly responsive, "
         "and secure foundation. By combining advanced natural language translation with professional institutional reporting "
         "and transaction rollback capabilities, the system reduces administrative workload, prevents data loss, and delivers "
         "an exceptionally smooth, premium user experience. The architectural designs and modules verified in this phase "
         "will serve as a solid launchpad for advanced AI integrations in future iterations.")
    ]
    
    for sec_title, sec_text in sections_content:
        # Chapter title
        p_ch = doc.add_paragraph()
        run_ch = p_ch.add_run(sec_title.upper())
        run_ch.font.size = Pt(13)
        run_ch.font.bold = True
        run_ch.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        p_ch.paragraph_format.space_before = Pt(12)
        p_ch.paragraph_format.space_after = Pt(6)
        
        # Chapter text
        p_body = doc.add_paragraph()
        p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_body.paragraph_format.line_spacing = 1.15
        p_body.paragraph_format.space_after = Pt(10)
        
        # Split text into bullet list if it has bullet marks or list structures
        lines = sec_text.split('\n')
        for line in lines:
            if line.strip().startswith("•") or line.strip().startswith("-") or (len(line.strip()) > 2 and line.strip()[1] == '.'):
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_after = Pt(4)
                run_bul = p_bullet.add_run(line.strip().lstrip('•-1234. ').strip())
                run_bul.font.size = Pt(11)
            elif line.strip().startswith("✔"):
                p_check = doc.add_paragraph()
                p_check.paragraph_format.left_indent = Inches(0.25)
                p_check.paragraph_format.space_after = Pt(4)
                run_chk = p_check.add_run(line.strip())
                run_chk.font.size = Pt(11)
                run_chk.font.bold = True
                run_chk.font.color.rgb = RGBColor(0x2A, 0x52, 0x98)
            else:
                if line.strip():
                    p_txt = doc.add_paragraph()
                    p_txt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p_txt.paragraph_format.space_after = Pt(8)
                    run_txt = p_txt.add_run(line.strip())
                    run_txt.font.size = Pt(11)
                    
        # Line break separator between chapters
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        
    doc.save(filename)
    print(f"[OK] DOCX report generated: {filename}")

def build_pdf(filename="Phase1_Project_Report.pdf"):
    print("Generating PDF...")
    doc = SimpleDocTemplate(
        filename, pagesize=A4,
        leftMargin=2.0*cm, rightMargin=2.0*cm,
        topMargin=2.0*cm, bottomMargin=2.0*cm
    )
    
    # Styles
    title_s = ParagraphStyle('T', fontName='Helvetica-Bold', fontSize=18, alignment=TA_CENTER, leading=22, textColor=colors.HexColor("#1e3a5f"))
    sub_s   = ParagraphStyle('S', fontName='Helvetica-Oblique', fontSize=12, alignment=TA_CENTER, leading=16, textColor=colors.HexColor("#555555"))
    univ_s  = ParagraphStyle('U', fontName='Helvetica-Bold', fontSize=13, alignment=TA_CENTER, leading=18, textColor=colors.HexColor("#1e3a5f"))
    dept_s  = ParagraphStyle('D', fontName='Helvetica-Bold', fontSize=10, alignment=TA_CENTER, leading=14, textColor=colors.HexColor("#2a5298"))
    
    ch_s    = ParagraphStyle('C', fontName='Helvetica-Bold', fontSize=12, alignment=TA_LEFT, leading=16, textColor=colors.HexColor("#1e3a5f"), spaceBefore=10, spaceAfter=5)
    body_s  = ParagraphStyle('B', fontName='Helvetica', fontSize=10, alignment=TA_JUSTIFY, leading=14, textColor=colors.HexColor("#333333"), spaceAfter=6)
    chk_s   = ParagraphStyle('CK', fontName='Helvetica-Bold', fontSize=9.5, alignment=TA_LEFT, leading=13, textColor=colors.HexColor("#2a5298"), leftIndent=15, spaceAfter=4)
    bul_s   = ParagraphStyle('BL', fontName='Helvetica', fontSize=10, alignment=TA_LEFT, leading=13, textColor=colors.HexColor("#333333"), leftIndent=15, spaceAfter=4)
    
    story = []
    
    # Title Page
    story.append(Spacer(1, 3*cm))
    story.append(Paragraph(UNIVERSITY.upper(), univ_s))
    story.append(Spacer(1, 4*cm))
    story.append(Paragraph(TITLE.upper(), title_s))
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph(SUBTITLE, sub_s))
    
    story.append(Spacer(1, 6*cm))
    story.append(Paragraph(DEPT, dept_s))
    story.append(Paragraph(COLLEGE, dept_s))
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph(DATE_STR, dept_s))
    story.append(PageBreak())
    
    # ── Report Content Definitions ─────────────────────────────────────────────
    sections_content = [
        ("1. Abstract", 
         "The AI & NLP Driven Smart Student Database Management System is a state-of-the-art enterprise-grade "
         "academic ERP system that bridges the gap between natural human language and structured relational databases. "
         "By leveraging advanced NLP techniques, fuzzy phonetic matching algorithms, and RAG (Retrieval-Augmented Generation), "
         "the system translates text and voice queries directly into secure, highly optimized MySQL commands. "
         "It addresses crucial bottlenecks in institutional operations, including complex search ambiguities, academic status tracking, "
         "robust schema migrations, and full transactional recovery with a multi-layered undo snapshot mechanism. "
         "The application delivers premium institutional branding, highly resilient data parsing, and absolute security against "
         "SQL injection attacks, redefining digital student administration in accordance with modern pedagogical guidelines."),
        
        ("2. Introduction",
         "In modern academic settings, administrative tasks demand speed, precision, and data integrity. Traditional ERP solutions "
         "rely on complex menu trees, multiple filter selections, and precise SQL matching, resulting in steep learning curves "
         "and operational delays. The 'AI & NLP Driven Smart Student Database Management System' is developed to resolve "
         "these limitations. By integrating intuitive search with relational MySQL databases, it allows users—both tech-savvy "
         "administrators and non-technical staff—to command the system via plain text and spoken language. "
         "This Phase-1 report outlines the complete architectural layout, database structures, AI-powered parsing engines, "
         "7-layer fuzzy matching modules, and robust soft-delete pipelines that make this system exceptionally resilient, user-friendly, "
         "and highly responsive."),
        
        ("3. Problem Statement",
         "Academic databases suffer from severe operational friction: \n"
         "1. Traditional search engines fail to find records when users make typos or phonetic naming errors, causing "
         "redundant entry or administrative delays.\n"
         "2. When multiple students share identical or similar names, typical systems automatically merge, overwrite, "
         "or arbitrarily display the first match without letting the user select the appropriate student, creating data contamination.\n"
         "3. Bulk file uploads of spreadsheets containing varied column layouts require manually rearranging spreadsheets "
         "to fit rigid database constraints, leading to massive manual parsing effort.\n"
         "4. Accidentally deleted or updated records are immediately permanent in traditional relational setups, lacking "
         "an instant point-in-time recovery without restoring bulky daily database backups."),
        
        ("4. Objectives",
         "The core objectives of the system include:\n"
         "• Develop a multi-layer NLP-to-SQL search pipeline utilizing token-level fuzzy, trigram, and phonetic matching "
         "to enable highly lenient live typing suggestions (≥0.20 score) and exact submitted queries (≥0.70 threshold).\n"
         "• Implement an ambiguity detection workflow that intercepts multi-student matches with similar names, blocking "
         "unauthorized merges and prompting the user to select the specific student by USN.\n"
         "• Build an AI Column Mapping and Wide-to-Long Excel Parser that auto-detects headers, matches custom columns to "
         "canonical database columns via exact/partial matching, and unpivots wide-format semester metrics to normalized rows.\n"
         "• Formulate a VTU Semester Progression Logic to automatically calculate estimated semesters, academic years, lateral entry "
         "variances, and graduated status based on parsed USN structures.\n"
         "• Incorporate an undo-snapshot engine offering a 5-minute transaction recovery window for deletion, updates, "
         "and bulk uploads by leveraging soft-deleted JSON snapshots."),
        
        ("5. Existing vs Proposed System",
         "The current project replaces outdated, rigid, and slow search architectures with a state-of-the-art system. "
         "The differences are highlighted below:\n\n"
         "• SEARCH LOGIC: The existing system utilizes rigid exact SQL substring matching (e.g., LIKE '%name%'), which fails on typos. "
         "The proposed system uses a 7-layer NLP pipeline incorporating soundex, double metaphone, trigram, and Levenshtein metrics "
         "to resolve spelling and phonetic discrepancies.\n"
         "• BULK INGESTION: Traditional systems crash or reject spreadsheets if column names do not exactly match the database. "
         "The proposed system utilizes a rule-based AI Column Mapper that normalizes column headers and automatically parses wide-format "
         "semester matrices to canonical, normalized structures.\n"
         "• DELETION SAFETY: In typical architectures, database deletions are immediately permanent, risking complete data loss. "
         "The proposed system captures full point-in-time pre-deletion JSON snapshots, allowing instant restore via a unique UUID "
         "token within a 5-minute window.\n"
         "• BRANDED EXPORTS: Output reporting in standard tools returns generic unformatted tables. The proposed system employs a unified "
         "branded report generator in PDF, Excel, and CSV formats, including institutional headers and computer-filled signature sections."),
        
        ("6. Technologies Used",
         "• Backend Framework: Python FastAPI (v0.109.2) - High-performance ASGI framework enabling concurrent async API execution.\n"
         "• Frontend Architecture: React.js (v18) & Vite - Lightning-fast UI component loading and responsive design using TailwindCSS.\n"
         "• Database Engine: Oracle MySQL (v8.0) - Relational DB utilizing indexing, unique keys, and transactions.\n"
         "• Text & Voice Parsing: Web Speech API - Browser-native voice dictation converting oral requests directly into strings.\n"
         "• Data Modeling & Math: Pandas (v2.2.0) - High-speed dataframe manipulation used in file parsing.\n"
         "• Professional Reporting: ReportLab (v4.1.0) & OpenPyXL (v3.1.2) - Direct generation of branded PDF, Excel, and CSV exports.\n"
         "• Security & Authentication: Bcrypt (v4.1.2) & Python-Jose (v3.3.0) - Robust password hashing and JWT token handling."),
        
        ("7. System Architecture & Working Pipeline",
         "The operation follows a clear, logical progression: \n"
         "1. USER QUERY: Plain text input is received via the React input terminal or Web Speech API.\n"
         "2. NLP PROCESSING: Stop words are eliminated; the query is analyzed for intent (personal, academic, or full) and search terms are extracted.\n"
         "3. SQL QUERY GENERATION: RAG (Retrieval-Augmented Generation) feeds schema constraints and few-shot examples to the NVIDIA LLM API.\n"
         "4. SECURITY VALIDATION: The generated SQL is parsed by the security validator. Unsafe commands (DROP, ALTER, TRUNCATE) or destructive statements lacking a WHERE clause are blocked and logged in the security_logs table.\n"
         "5. DATABASE FETCH: Valid SELECT queries are fetched from the MySQL connection pool. If 0 records are returned, the system falls back to phonetic/fuzzy matching.\n"
         "6. DUAL-MODAL FEEDBACK: Ambiguities (multi-student name matches) are displayed via selection cards. Confirmed results are mapped and rendered in interactive tables, charts, or exported in high-quality formats."),
        
        ("8. Modules Description",
         "• NLP Parsing & Intent Detection: Strips stop words, determines query scope (academic vs. personal), and extracts student names/USNs.\n"
         "• 7-Layer Fuzzy Search Engine: Provides live, lenient, spelling-resilient suggestions during typing and strict confidence scoring upon submission.\n"
         "• AI File Upload & Column Mapper: Receives CSV/Excel files, normalizes columns, and unpivots wide semester columns into structured rows.\n"
         "• Unified Export System: Builds identical Report models to construct branded PDFs, Excel sheets, and CSVs with integrated signatures.\n"
         "• Undo/Snapshot Engine: Intercepts updates, deletes, and uploads to save pre-state JSON dumps for instant Point-In-Time rollback."),
        
        ("9. 7-Layer NLP Search Engine Logic",
         "Spelling-resilient matching is achieved using a robust 7-layer hybrid scoring and search pipeline:\n"
         "1. Exact String Match: Direct case-sensitive string equality checks (100% confidence).\n"
         "2. Normalized Exact Match: Lowercases input, strips accents via unicodedata, and trims special characters (97% confidence).\n"
         "3. Prefix & Token Prefix Match: Checks if candidate names or space-separated name tokens start with the query (e.g. 'rut' -> 'ruthik'). Utilizes a length-based coverage score ratio.\n"
         "4. Substring & Token Substring Match: Checks if normalized query exists inside candidate name tokens (e.g. 'anth' -> 'manjunath').\n"
         "5. Phonetic Encoding (Soundex & Double Metaphone): Generates standard 4-char English pronunciation keys (Soundex) and advanced multi-pronunciation hashes (Double Metaphone) via Jellyfish to identify words sounding identical despite spelling differences (e.g., 'Sudheer' -> 'Sudhir').\n"
         "6. Character Trigram Similarity: Extracts character-level trigrams and computes a similarity coefficient matching typos: 2.0 * len(intersection) / (len(trigrams_A) + len(trigrams_B)).\n"
         "7. Levenshtein & Jaro-Winkler String Distance: Calculates single-character insertions, deletions, substitutions, and transpositions to grade typographical proximity.\n"
         "• Live Suggestion Rule: Executes leniently with score threshold >= 0.20 for high-response keyboard inputs.\n"
         "• Submission Execution Rule: Applies strict score threshold >= 0.70 to ensure secure database write-back commands."),
        
        ("10. Smart Semester Progression Logic",
         "The system parses Visvesvaraya Technological University (VTU) USN structures (e.g. '1GC20CS042') using regex:\n"
         "• Course Duration & Year: First character digit defines course duration (e.g., 4 years). Characters 4-5 represent the short admission year (e.g. 20 -> 2020).\n"
         "• Student Type: Roll number digits (characters 8-10) are parsed to detect Lateral Entry (roll number >= 400). Regular students start from 1.\n"
         "• Progression Calculation: Compares admission year to current year. If the current month is >= July, the semester is calculated as (years_diff * 2 + 1) [Odd Semester]. Otherwise, it is (years_diff * 2) [Even Semester]. Lateral entry students automatically receive a +2 semester adjustment.\n"
         "• Status Mapping: If the estimated semester exceeds (course_duration * 2), status is set to GRADUATED; otherwise, ACTIVE."),
        
        ("11. File Ingestion & AI Mapping",
         "Spreadsheet upload is managed through a strict, zero-loss pipeline:\n"
         "1. Header Row Detection: Scans the first 20 rows of Excel/CSV sheets to locate the real header by identifying known keywords like USN, name, or GPA.\n"
         "2. AI Column Mapper: Maps custom headers to canonical database columns. For example, 'Reg Number' maps to 'usn', and 'Fathers Name' to 'father_name'.\n"
         "3. Wide-Format Unpivoting: Converts wide sheets (columns: sem_1_sgpa, sem_2_sgpa) into standardized normalized rows (Semester, SGPA).\n"
         "4. USN Validation & Merge: Checks USN structure using standard regex. If valid, records are merged (upserted); duplicate rows are identified and flagged."),
        
        ("12. Undo & Point-in-Time Recovery",
         "To guarantee complete transactional safety, a 5-minute Point-In-Time recovery window is enforced:\n"
         "• Deletion/Update Interception: Prior to executing any destructive command, the system runs a select query to fetch all target rows.\n"
         "• Snapshot Generation: Stores the full student profiles and academic marks as a serialized JSON dump in the `global_undo_snapshots` table, returning a unique UUID token.\n"
         "• Recovery Execution: Upon post-requesting the token to `/api/undo/restore/{token}`, the system deletes the newly modified data, reads the JSON snapshot, and re-inserts the exact original values."),
        
        ("13. MySQL Database Design",
         "The database is designed with optimized indexing and strict integrity constraints. Key tables include:\n"
         "• users: Manages credentials, roles (Admin/Staff), and preferences (theme) with custom ENUMs.\n"
         "• students: Central registry of personal fields (USN, name, DOB, father name, permanent address) with a UNIQUE USN constraint.\n"
         "• marks: Tracks academic data (SGPA, CGPA) linked to USN via foreign key cascading, protected by a unique index (uq_usn_sem).\n"
         "• global_undo_snapshots: Stores the pre-operation JSON data, operation type, actor, and status for point-in-time rollbacks.\n"
         "• query_history & security_logs: Maintain detailed auditing for natural queries, executed SQL statements, and intercepted injection attempts."),
        
        ("14. UI/UX Architecture",
         "The frontend is constructed using modular React component design:\n"
         "• Dashboard Workspace: Provides a single-page reactive shell hosting a collapsible sidebar and unified search hub.\n"
         "• Live Suggestion Panel: A highly interactive dropdown showing spelling matches with badges (EXACT, FUZZY, PHONETIC).\n"
         "• Interactive Data Tables: Generates responsive paginated grids with column sorting and search filtering.\n"
         "• Activity Feed & Rollback Panel: Allows administrators to review security logs and trigger instant undo rollbacks in real time."),
        
        ("15. Results & Implemented Features",
         "The Phase-1 implementation succeeds in delivering all core modules:\n"
         "✔ Verified 7-layer fuzzy phonetic search engine resolving severe typos.\n"
         "✔ Operational async text and voice-dictated query generation pipeline.\n"
         "✔ Zero-loss spreadsheet ingestion engine with automatic unpivoting.\n"
         "✔ Strict security validation blocking malicious AST injection patterns.\n"
         "✔ Verified 5-minute Point-In-Time Undo Recovery with JSON snapshots.\n"
         "✔ Institutional branded PDF, Excel, and CSV export modules."),
        
        ("16. Remaining Issues & Future Scope",
         "Current limitations and developmental plans for Phase-2 include:\n"
         "• Current Limitations: Local MySQL service configuration dependencies (requires manual administrative startup); single-instance file upload cache (susceptible to cache expiration on server restart).\n"
         "• Future Scope: Integration of vector embedding semantic search using Milvus/ChromaDB to handle complex intent mappings; implementation of Redis-based query caching to bypass LLM generation for identical queries; integration of automated WebSockets for multi-user real-time notification feeds."),
        
        ("17. Conclusion",
         "Phase-1 of the AI & NLP Driven Student Database Management System successfully establishes a robust, highly responsive, "
         "and secure foundation. By combining advanced natural language translation with professional institutional reporting "
         "and transaction rollback capabilities, the system reduces administrative workload, prevents data loss, and delivers "
         "an exceptionally smooth, premium user experience. The architectural designs and modules verified in this phase "
         "will serve as a solid launchpad for advanced AI integrations in future iterations.")
    ]
    
    for sec_title, sec_text in sections_content:
        # Section Heading
        story.append(Paragraph(sec_title.upper(), ch_s))
        story.append(HRFlowable(width="100%", thickness=0.8, color=colors.HexColor("#1e3a5f"), spaceAfter=6))
        
        lines = sec_text.split('\n')
        for line in lines:
            if line.strip().startswith("•") or line.strip().startswith("-") or (len(line.strip()) > 2 and line.strip()[1] == '.'):
                story.append(Paragraph(line.strip(), bul_s))
            elif line.strip().startswith("✔"):
                story.append(Paragraph(line.strip(), chk_s))
            else:
                if line.strip():
                    story.append(Paragraph(line.strip(), body_s))
        story.append(Spacer(1, 0.4*cm))
        
    doc.build(story)
    print(f"[OK] PDF report generated: {filename}")

if __name__ == "__main__":
    build_docx()
    build_pdf()
